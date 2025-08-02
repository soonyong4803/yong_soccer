pip install python-docx
python - <<'PY'
from docx import Document

doc = Document()
doc.add_paragraph("== [20250802-K2-GYE-BUS-GYE] ==")
doc.add_paragraph("qual_injury: 핵심 부상 없음")
doc.add_paragraph("qual_lineup: 외국인 선발 예상")
doc.add_paragraph("qual_tactics: 역습 기조 예상")
doc.add_paragraph("qual_motivation: 홈 반등 필요")
doc.add_paragraph("qual_weather: 무더위 변수")
doc.add_paragraph("")
doc.add_paragraph("== [20250802-K2-GYE-BUS-BUS] ==")
doc.add_paragraph("qual_injury: 중원 리더 결장")
doc.add_paragraph("qual_lineup: 최근 공격진 변경")
doc.add_paragraph("qual_tactics: 강한 압박 전술 유지")
doc.add_paragraph("qual_motivation: 리벤지 의지 높음")
doc.add_paragraph("qual_weather: 적응된 여건")
doc.save("qual_multigame_template.docx")
print("✅ qual_multigame_template.docx 생성 완료")
PY
